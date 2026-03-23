from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT_DIR = Path(__file__).resolve().parents[4]
DEFAULT_SOURCE = ROOT_DIR / 'docs' / 'article' / 'eduassist-platform-academic-article.md'
DEFAULT_METADATA = ROOT_DIR / 'docs' / 'article' / 'article-export-metadata.yaml'
DEFAULT_OUTPUT = ROOT_DIR / 'artifacts' / 'article' / 'eduassist-platform-academic-article.docx'


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='Export the academic article to DOCX.')
    parser.add_argument('--source', type=Path, default=DEFAULT_SOURCE)
    parser.add_argument('--metadata', type=Path, default=DEFAULT_METADATA)
    parser.add_argument('--output', type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def _load_metadata(path: Path) -> dict[str, Any]:
    return yaml.safe_load(path.read_text(encoding='utf-8')) or {}


def _set_font(style: Any, *, name: str, size: int, bold: bool | None = None) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def _configure_styles(document: Document) -> None:
    _set_font(document.styles['Normal'], name='Times New Roman', size=12)
    _set_font(document.styles['Title'], name='Times New Roman', size=16, bold=True)
    _set_font(document.styles['Heading 1'], name='Times New Roman', size=14, bold=True)
    _set_font(document.styles['Heading 2'], name='Times New Roman', size=12, bold=True)
    _set_font(document.styles['Heading 3'], name='Times New Roman', size=12, bold=True)
    if 'List Bullet' in document.styles:
        _set_font(document.styles['List Bullet'], name='Times New Roman', size=12)
    if 'List Number' in document.styles:
        _set_font(document.styles['List Number'], name='Times New Roman', size=12)


def _clean_inline(text: str) -> str:
    text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', text)
    text = text.replace('**', '')
    text = text.replace('`', '')
    return text.strip()


def _add_cover(document: Document, metadata: dict[str, Any]) -> None:
    title = document.add_paragraph(style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(str(metadata.get('title', 'Titulo a definir')))

    subtitle_value = str(metadata.get('subtitle', '')).strip()
    if subtitle_value:
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(subtitle_value)
        subtitle_run.italic = True
        subtitle_run.font.name = 'Times New Roman'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        subtitle_run.font.size = Pt(12)

    document.add_paragraph('')

    for key in ('author', 'institution', 'program', 'advisor'):
        value = str(metadata.get(key, '')).strip()
        if not value:
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(value)

    document.add_paragraph('')

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    location = str(metadata.get('location', '')).strip()
    date = str(metadata.get('date', '')).strip()
    footer.add_run(' - '.join(part for part in (location, date) if part))
    document.add_page_break()


def _render_markdown(document: Document, markdown_text: str) -> None:
    buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal buffer
        if not buffer:
            return
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.add_run(_clean_inline(' '.join(buffer)))
        buffer = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith('# '):
            continue

        if not stripped:
            flush_paragraph()
            continue

        if stripped.startswith('## '):
            flush_paragraph()
            document.add_heading(_clean_inline(stripped[3:]), level=1)
            continue

        if stripped.startswith('### '):
            flush_paragraph()
            document.add_heading(_clean_inline(stripped[4:]), level=2)
            continue

        if re.match(r'^\d+\.\s+', stripped):
            flush_paragraph()
            item = re.sub(r'^\d+\.\s+', '', stripped)
            document.add_paragraph(_clean_inline(item), style='List Number')
            continue

        if stripped.startswith('- '):
            flush_paragraph()
            document.add_paragraph(_clean_inline(stripped[2:]), style='List Bullet')
            continue

        buffer.append(stripped)

    flush_paragraph()


def main() -> int:
    args = parse_args()
    metadata = _load_metadata(args.metadata.resolve())
    source_text = args.source.resolve().read_text(encoding='utf-8')
    output_path = args.output.resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_styles(document)
    document.core_properties.title = str(metadata.get('title', 'EduAssist Platform'))
    document.core_properties.author = str(metadata.get('author', 'Autor a definir'))
    document.core_properties.subject = 'Artigo academico do projeto EduAssist Platform'

    _add_cover(document, metadata)
    _render_markdown(document, source_text)
    document.save(output_path)

    print(
        {
            'source': args.source.resolve().as_posix(),
            'metadata': args.metadata.resolve().as_posix(),
            'output': output_path.as_posix(),
        }
    )
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
